"""Crawl approved public sources into a local SQLite document store.

The crawler deliberately follows only URLs listed in sources.json.  Add each
new NJUPT subsite explicitly after reviewing its public-access rules.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sqlite3
import sys
import time
from collections import deque
from dataclasses import dataclass
from datetime import UTC, datetime
from html.parser import HTMLParser
from http.client import IncompleteRead
from io import BytesIO
from pathlib import Path
from typing import Iterable
from urllib.error import HTTPError, URLError
from urllib.parse import unquote, urljoin, urlsplit, urlunsplit
from urllib.request import Request, urlopen


USER_AGENT = "NJUPT-Public-RAG/0.1 (allowlist crawler)"
PROJECT_ROOT = Path(__file__).resolve().parents[1]
DATE_PATTERN = re.compile(r"(20\d{2})[年./-]\s*(\d{1,2})[月./-]\s*(\d{1,2})")
ATTACHMENT_SUFFIXES = (".pdf", ".docx")
MAX_DOWNLOAD_BYTES = 128 * 1024 * 1024
DOWNLOAD_ATTEMPTS = 3


@dataclass(frozen=True)
class Source:
    id: str
    category: str
    start_urls: tuple[str, ...]
    allowed_hosts: tuple[str, ...]
    url_patterns: tuple[re.Pattern[str], ...]
    excluded_patterns: tuple[re.Pattern[str], ...]


class PageParser(HTMLParser):
    """Small dependency-free text and link extractor for public HTML pages."""

    ignored_tags = {"script", "style", "noscript", "svg", "form", "iframe"}
    boundary_tags = {"p", "div", "li", "h1", "h2", "h3", "h4", "tr", "br"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.title_parts: list[str] = []
        self.text_parts: list[str] = []
        self.links: list[str] = []
        self._ignore_depth = 0
        self._in_title = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in self.ignored_tags:
            self._ignore_depth += 1
        if tag == "title":
            self._in_title = True
        if tag == "a":
            href = dict(attrs).get("href")
            if href:
                self.links.append(href)
        if tag in self.boundary_tags:
            self.text_parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self.ignored_tags and self._ignore_depth:
            self._ignore_depth -= 1
        if tag == "title":
            self._in_title = False
        if tag in self.boundary_tags:
            self.text_parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._ignore_depth:
            return
        text = " ".join(data.split())
        if not text:
            return
        if self._in_title:
            self.title_parts.append(text)
        self.text_parts.append(text)

    @property
    def title(self) -> str:
        return " ".join(self.title_parts).strip()

    @property
    def text(self) -> str:
        lines = (" ".join(line.split()) for line in "".join(self.text_parts).splitlines())
        return "\n".join(line for line in lines if line)


def load_sources(path: Path) -> list[Source]:
    raw = json.loads(path.read_text(encoding="utf-8"))
    return [
        Source(
            id=item["id"],
            category=item["category"],
            start_urls=tuple(item["start_urls"]),
            allowed_hosts=tuple(item["allowed_hosts"]),
            url_patterns=tuple(re.compile(pattern) for pattern in item["url_patterns"]),
            excluded_patterns=tuple(re.compile(pattern, re.IGNORECASE) for pattern in item["excluded_patterns"]),
        )
        for item in raw["sources"]
    ]


def normalize_url(url: str) -> str:
    parts = urlsplit(url)
    return urlunsplit((parts.scheme.lower(), parts.netloc.lower(), parts.path or "/", parts.query, ""))


def is_allowed(url: str, source: Source) -> bool:
    parts = urlsplit(url)
    if parts.scheme not in {"http", "https"} or parts.hostname not in source.allowed_hosts:
        return False
    candidate = f"{parts.hostname}{parts.path}?{parts.query}"
    if any(pattern.search(candidate) for pattern in source.excluded_patterns):
        return False
    return any(pattern.search(parts.path) for pattern in source.url_patterns) or parts.path.lower().endswith(ATTACHMENT_SUFFIXES)


def extract_published_at(text: str) -> str | None:
    match = DATE_PATTERN.search(text)
    if not match:
        return None
    year, month, day = (int(value) for value in match.groups())
    try:
        return datetime(year, month, day).date().isoformat()
    except ValueError:
        return None


def extract_pdf(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF, managed by uv in pyproject.toml
    except ImportError as exc:
        raise RuntimeError("PDF text extraction needs PyMuPDF: uv add PyMuPDF") from exc
    document = fitz.open(stream=data, filetype="pdf")
    return "\n".join(page.get_text() for page in document).strip()


def extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX text extraction needs python-docx: uv add python-docx") from exc
    document = Document(BytesIO(data))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def detect_source_type(url: str, content_type: str) -> str:
    suffix = Path(urlsplit(url).path).suffix.lower()
    if suffix in ATTACHMENT_SUFFIXES:
        return suffix.removeprefix(".")
    content_types = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    }
    if content_type in content_types:
        return content_types[content_type]
    if content_type in {"text/html", "application/xhtml+xml"}:
        return "html"
    raise RuntimeError(f"unsupported content type: {content_type}")


class Store:
    def __init__(self, db_path: Path, raw_dir: Path) -> None:
        db_path.parent.mkdir(parents=True, exist_ok=True)
        raw_dir.mkdir(parents=True, exist_ok=True)
        self.raw_dir = raw_dir
        self.connection = sqlite3.connect(db_path)
        self.connection.execute(
            """CREATE TABLE IF NOT EXISTS documents (
                doc_id TEXT PRIMARY KEY, url TEXT UNIQUE NOT NULL, title TEXT,
                text TEXT NOT NULL, category TEXT NOT NULL, published_at TEXT,
                crawled_at TEXT NOT NULL, content_hash TEXT NOT NULL,
                source_type TEXT NOT NULL, raw_path TEXT NOT NULL
            )"""
        )
        self.connection.execute(
            """CREATE TABLE IF NOT EXISTS crawl_failures (
                url TEXT PRIMARY KEY, source_id TEXT NOT NULL, error TEXT NOT NULL,
                failed_at TEXT NOT NULL
            )"""
        )
        self.connection.commit()

    def save_document(
        self, url: str, title: str, text: str, source: Source, published_at: str | None,
        source_type: str, raw: bytes, crawled_at: str,
    ) -> None:
        content_hash = hashlib.sha256(raw).hexdigest()
        suffix = f".{source_type}"
        raw_path = self.raw_dir / f"{content_hash}{suffix}"
        if not raw_path.exists():
            raw_path.write_bytes(raw)
        doc_id = hashlib.sha256(url.encode("utf-8")).hexdigest()
        self.connection.execute(
            """INSERT INTO documents
            (doc_id, url, title, text, category, published_at, crawled_at, content_hash, source_type, raw_path)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(url) DO UPDATE SET title=excluded.title, text=excluded.text,
                category=excluded.category, published_at=excluded.published_at,
                crawled_at=excluded.crawled_at, content_hash=excluded.content_hash,
                source_type=excluded.source_type, raw_path=excluded.raw_path""",
            (doc_id, url, title, text, source.category, published_at, crawled_at, content_hash, source_type, str(raw_path)),
        )
        self.connection.execute("DELETE FROM crawl_failures WHERE url = ?", (url,))
        self.connection.commit()

    def save_failure(self, url: str, source: Source, error: str) -> None:
        self.connection.execute(
            """INSERT INTO crawl_failures (url, source_id, error, failed_at) VALUES (?, ?, ?, ?)
            ON CONFLICT(url) DO UPDATE SET source_id=excluded.source_id, error=excluded.error, failed_at=excluded.failed_at""",
            (url, source.id, error, utc_now()),
        )
        self.connection.commit()

    def close(self) -> None:
        self.connection.close()


def utc_now() -> str:
    return datetime.now(UTC).isoformat(timespec="seconds")


def download(url: str) -> tuple[bytes, str]:
    request = Request(url, headers={"User-Agent": USER_AGENT, "Accept": "text/html,application/pdf;q=0.9,*/*;q=0.1"})
    for attempt in range(DOWNLOAD_ATTEMPTS):
        try:
            with urlopen(request, timeout=20) as response:
                declared_size = response.headers.get("Content-Length")
                if declared_size and int(declared_size) > MAX_DOWNLOAD_BYTES:
                    raise RuntimeError(f"file exceeds {MAX_DOWNLOAD_BYTES // 1024 // 1024} MiB limit")
                chunks: list[bytes] = []
                size = 0
                while chunk := response.read(1024 * 1024):
                    size += len(chunk)
                    if size > MAX_DOWNLOAD_BYTES:
                        raise RuntimeError(f"file exceeds {MAX_DOWNLOAD_BYTES // 1024 // 1024} MiB limit")
                    chunks.append(chunk)
                raw = b"".join(chunks)
                if declared_size and len(raw) != int(declared_size):
                    raise IncompleteRead(raw, int(declared_size) - len(raw))
                return raw, response.headers.get_content_type()
        except HTTPError:
            raise
        except (IncompleteRead, URLError, TimeoutError, OSError) as exc:
            if attempt == DOWNLOAD_ATTEMPTS - 1:
                raise RuntimeError(f"download failed after {DOWNLOAD_ATTEMPTS} attempts: {exc}") from exc
            time.sleep(2 ** attempt)
    raise AssertionError("unreachable")


def crawl_source(source: Source, store: Store, max_pages: int, delay: float, dry_run: bool) -> tuple[int, int]:
    queue = deque(normalize_url(url) for url in source.start_urls)
    visited: set[str] = set()
    succeeded = failed = 0
    while queue and len(visited) < max_pages:
        url = queue.popleft()
        if url in visited or not is_allowed(url, source):
            continue
        visited.add(url)
        try:
            raw, content_type = download(url)
            crawled_at = utc_now()
            source_type = detect_source_type(url, content_type)
            if source_type == "pdf":
                text = extract_pdf(raw)
                links: Iterable[str] = ()
            elif source_type == "docx":
                text = extract_docx(raw)
                links = ()
            else:
                parser = PageParser()
                parser.feed(raw.decode("utf-8", errors="replace"))
                text, title, links = parser.text, parser.title, parser.links
            if source_type in {"pdf", "docx"}:
                title = unquote(Path(urlsplit(url).path).name)
            if not text:
                raise RuntimeError("empty extracted text")
            if dry_run:
                print(f"DRY-RUN {source.id}: {url}")
            else:
                store.save_document(url, title, text, source, extract_published_at(text), source_type, raw, crawled_at)
                print(f"SAVED {source.id}: {url}")
            succeeded += 1
            for link in links:
                candidate = normalize_url(urljoin(url, link))
                if candidate not in visited and is_allowed(candidate, source):
                    queue.append(candidate)
        except (HTTPError, URLError, TimeoutError, RuntimeError, ValueError) as exc:
            failed += 1
            print(f"FAILED {source.id}: {url} ({exc})", file=sys.stderr)
            if not dry_run:
                store.save_failure(url, source, str(exc))
        if queue and delay:
            time.sleep(delay)
    return succeeded, failed


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--config", type=Path, default=Path(__file__).with_name("sources.json"))
    parser.add_argument("--database", type=Path, default=PROJECT_ROOT / "data" / "crawler.db")
    parser.add_argument("--raw-dir", type=Path, default=PROJECT_ROOT / "data" / "raw")
    parser.add_argument("--source", action="append", help="Source id to crawl; repeat to select several")
    parser.add_argument("--max-pages", type=int, default=10, help="Maximum pages per source in this run")
    parser.add_argument("--delay", type=float, default=1.5, help="Seconds to wait between requests")
    parser.add_argument("--dry-run", action="store_true", help="Fetch and inspect without writing SQLite/raw files")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if args.max_pages < 1 or args.delay < 0:
        raise SystemExit("--max-pages must be positive and --delay cannot be negative")
    selected = [source for source in load_sources(args.config) if not args.source or source.id in args.source]
    if args.source and len(selected) != len(set(args.source)):
        raise SystemExit("one or more --source values do not exist in the configuration")
    store = Store(args.database, args.raw_dir)
    try:
        for source in selected:
            succeeded, failed = crawl_source(source, store, args.max_pages, args.delay, args.dry_run)
            print(f"{source.id}: {succeeded} saved, {failed} failed")
    finally:
        store.close()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
